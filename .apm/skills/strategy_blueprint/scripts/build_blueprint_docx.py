"""Build a KaxaNuk Strategy Blueprint Word document from a content.json file.

This is the Word counterpart of ``build_blueprint.py`` (PDF). It reads the
exact same ``content.json`` schema (see ``references/content_schema.md``) and
produces a .docx that mirrors the same KaxaNuk brand: orange section
headings, Montserrat throughout, the small KaxaNuk mark in the page header
of every page after the cover, footer page numbers, and a colour-coded
Strategy Score table.

Run from the skill folder root:

    python scripts/build_blueprint_docx.py \
        --content content.json \
        --strategy-name "Liquidity-Weighted Trend Strategy" \
        --output Liquidity_Weighted_Trend_Strategy_Blueprint.docx

Dependencies:
    pip install python-docx pillow
"""
from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path
from typing import Any, Optional

try:
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.table import _Cell
except ImportError:  # pragma: no cover
    print(
        "python-docx is not installed. Run: pip install python-docx pillow",
        file=sys.stderr,
    )
    raise

# ---------------------------------------------------------------------------
# Brand constants (mirror references/BRANDING_GUIDELINES.md)
# ---------------------------------------------------------------------------

ORANGE = RGBColor(0xE8, 0x43, 0x28)
BODY = RGBColor(0x00, 0x00, 0x00)
BODY_SOFT = RGBColor(0x43, 0x43, 0x43)
FOOTER_GREY = RGBColor(0x7F, 0x7F, 0x7F)
DIVIDER_HEX = "C7C7C7"
LIGHT_BORDER_HEX = "E1E1E1"
SCORE_HIGH_HEX = "008A0E"
SCORE_MID_HEX = "FFE599"
SCORE_LOW_HEX = "F8696B"
WHITE_HEX = "FFFFFF"
LINK_BLUE_HEX = "1071E5"

FONT_REG = "Montserrat"
FONT_BOLD = "Montserrat"  # python-docx applies bold via run.bold

# Heading sizes (pt) per BRANDING_GUIDELINES.md
SIZE_COVER_TITLE = 36
SIZE_COVER_SUBTITLE = 20
SIZE_H1 = 24
SIZE_H2 = 20
SIZE_H3 = 16
SIZE_H4 = 14
SIZE_BODY = 10
SIZE_SIGNAL = 9

ASSETS = Path(__file__).resolve().parent.parent / "assets"
HEADER_MARK = ASSETS / "kaxanuk_mark_header.png"
COVER_LOGO = ASSETS / "kaxanuk_logo_full.png"


# ---------------------------------------------------------------------------
# Low-level helpers
# ---------------------------------------------------------------------------

def _set_run(
    run, *, size: int, bold: bool = False, color: RGBColor = BODY,
    font: str = FONT_REG,
) -> None:
    run.font.name = font
    # Set East Asian font too so Word doesn't substitute
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def _add_paragraph(doc_or_cell, text: str = "", *, size: int = SIZE_BODY,
                   bold: bool = False, color: RGBColor = BODY,
                   align: int = WD_ALIGN_PARAGRAPH.LEFT,
                   space_before: int = 0, space_after: int = 6,
                   left_indent: float = 0):
    para = doc_or_cell.add_paragraph()
    para.alignment = align
    pf = para.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if left_indent:
        pf.left_indent = Inches(left_indent)
    if text:
        run = para.add_run(text)
        _set_run(run, size=size, bold=bold, color=color)
    return para


def _add_heading(doc, text: str, level: int) -> None:
    """Render an H1-H4 heading.

    We deliberately do NOT apply Word's built-in ``Heading 1``/``Heading 2``
    styles. Those styles carry a theme colour and several Office-suite
    converters (notably LibreOffice's PDF export) treat the *style* colour as
    authoritative, silently replacing the brand orange even when the run sets
    its own colour. Styling the run directly is the only reliable way to
    keep ``#E84328`` on every heading regardless of which tool opens the
    document.
    """
    spec = {
        1: (SIZE_H1, ORANGE, 24, 8),
        2: (SIZE_H2, ORANGE, 18, 6),
        3: (SIZE_H3, ORANGE, 12, 4),
        4: (SIZE_H4, BODY,   10, 4),
    }[level]
    size, color, before, after = spec
    para = doc.add_paragraph()
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.keep_with_next = True
    run = para.add_run(text)
    _set_run(run, size=size, bold=True, color=color)


def _add_bullet(doc_or_cell, text: str, *, level: int = 0) -> None:
    """Render bullet using the literal ``●`` glyph to match the source PDF."""
    indent = 0.25 + 0.25 * level
    para = doc_or_cell.add_paragraph()
    para.paragraph_format.left_indent = Inches(indent)
    para.paragraph_format.space_after = Pt(2)
    para.paragraph_format.first_line_indent = Inches(-0.18)
    run = para.add_run("● ")
    _set_run(run, size=SIZE_BODY)
    run2 = para.add_run(text)
    _set_run(run2, size=SIZE_BODY)


def _add_arrow(doc, text: str) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.25)
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run("➔ ")
    _set_run(run, size=SIZE_BODY)
    run2 = para.add_run(text)
    _set_run(run2, size=SIZE_BODY)


def _shade(cell: _Cell, hex_color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _border(cell: _Cell, sides: tuple = ("top", "left", "bottom", "right"),
            *, sz: int = 4, color: str = DIVIDER_HEX) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for side in sides:
        tag = f"w:{side}"
        elem = tcBorders.find(qn(tag))
        if elem is None:
            elem = OxmlElement(tag)
            tcBorders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), str(sz))
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), color)


def _cell_text(cell: _Cell, text: str, *, size: int = SIZE_BODY,
               bold: bool = False, color: RGBColor = BODY,
               align: int = WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.text = ""  # python-docx leaves an empty para we'll reuse
    para = cell.paragraphs[0]
    para.alignment = align
    run = para.add_run(text)
    _set_run(run, size=size, bold=bold, color=color)


def _interpolate(value: float, lo: float, hi: float) -> str:
    """Hex string for the 3-color gradient SCORE_LOW -> SCORE_MID -> SCORE_HIGH."""
    if hi == lo:
        return SCORE_MID_HEX
    t = max(0.0, min(1.0, (value - lo) / (hi - lo)))
    if t < 0.5:
        s = t / 0.5
        rgb1 = (0xF8, 0x69, 0x6B)
        rgb2 = (0xFF, 0xE5, 0x99)
    else:
        s = (t - 0.5) / 0.5
        rgb1 = (0xFF, 0xE5, 0x99)
        rgb2 = (0x00, 0x8A, 0x0E)
    r = int(rgb1[0] + (rgb2[0] - rgb1[0]) * s)
    g = int(rgb1[1] + (rgb2[1] - rgb1[1]) * s)
    b = int(rgb1[2] + (rgb2[2] - rgb1[2]) * s)
    return f"{r:02X}{g:02X}{b:02X}"


def _set_header_with_mark(section) -> None:
    """Insert the KaxaNuk circular mark at the top-right of the page header."""
    header = section.header
    header.is_linked_to_previous = False
    para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if HEADER_MARK.exists():
        run = para.add_run()
        try:
            run.add_picture(str(HEADER_MARK), height=Inches(0.45))
        except Exception as exc:  # pragma: no cover
            print(f"WARN: could not embed header mark: {exc}", file=sys.stderr)
    else:
        print(f"WARN: header mark missing at {HEADER_MARK}", file=sys.stderr)


def _add_page_number_footer(section) -> None:
    footer = section.footer
    footer.is_linked_to_previous = False
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    _set_run(run, size=SIZE_BODY, color=FOOTER_GREY)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def _clear_header_footer(section) -> None:
    section.different_first_page_header_footer = False
    # Just leave them empty; for the cover we use a separate section.


def _page_break(doc) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


# ---------------------------------------------------------------------------
# Cover and structural pages
# ---------------------------------------------------------------------------

def render_cover(doc, content: dict) -> None:
    section = doc.sections[0]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    # Cover has no header/footer
    _clear_header_footer(section)

    _add_paragraph(doc, "", space_after=0)
    for _ in range(6):
        _add_paragraph(doc, "", space_after=0)
    _add_paragraph(doc, "Strategy Blueprint",
                   size=SIZE_COVER_TITLE, color=ORANGE,
                   align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
    for _ in range(4):
        _add_paragraph(doc, "", space_after=0)

    if COVER_LOGO.exists():
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        try:
            run.add_picture(str(COVER_LOGO), width=Inches(3.0))
        except Exception as exc:
            print(f"WARN: cover logo failed: {exc}", file=sys.stderr)
    else:
        print(f"WARN: cover logo missing at {COVER_LOGO}", file=sys.stderr)

    for _ in range(8):
        _add_paragraph(doc, "", space_after=0)
    _add_paragraph(doc, content.get("strategy_name", ""),
                   size=SIZE_COVER_SUBTITLE, bold=True,
                   align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)


def _begin_content_section(doc) -> None:
    """Start a new section so subsequent pages get the brand header/footer."""
    new_section = doc.add_section(WD_SECTION.NEW_PAGE)
    new_section.page_height = Inches(11)
    new_section.page_width = Inches(8.5)
    new_section.top_margin = Inches(1.0)
    new_section.bottom_margin = Inches(1.0)
    new_section.left_margin = Inches(1.0)
    new_section.right_margin = Inches(1.0)
    _set_header_with_mark(new_section)
    _add_page_number_footer(new_section)


def render_process_flow(doc, pf: dict) -> None:
    # Process Flow lives on its own page in a new section that carries the
    # KaxaNuk header mark and the page-number footer. The cover section that
    # precedes it stays clean.
    _begin_content_section(doc)
    _add_heading(doc, "Table of Contents", 1)
    _add_paragraph(doc, pf.get("intro", "The Process Flow:"),
                   size=15, bold=True, space_after=10)
    for item in pf.get("items", []):
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(10)
        run = para.add_run(item.get("label", ""))
        _set_run(run, size=15, bold=True)
        run2 = para.add_run(f" - {item.get('description', '')}")
        _set_run(run2, size=15)


# ---------------------------------------------------------------------------
# Main section renderers
# ---------------------------------------------------------------------------

def _image_or_placeholder(doc, image_path: Optional[str], caption: str) -> None:
    full = Path(image_path) if image_path else None
    if full and full.exists():
        try:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run()
            run.add_picture(str(full), width=Inches(6.0))
        except Exception as exc:
            print(f"WARN: could not embed {full}: {exc}", file=sys.stderr)
            _add_paragraph(doc, f"[{caption} pending]",
                           color=BODY_SOFT, align=WD_ALIGN_PARAGRAPH.CENTER)
    else:
        _add_paragraph(doc, f"[{caption} pending]",
                       color=BODY_SOFT, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(doc, caption, color=BODY_SOFT,
                   align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)


def _generic_table(doc, headers: list[str], rows: list[list[str]],
                   col_widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    if col_widths is None:
        col_widths = [6.5 / len(headers)] * len(headers)
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        _cell_text(cell, h, size=SIZE_BODY, bold=True, color=BODY)
        cell.width = Inches(col_widths[j])
        _border(cell, sides=("bottom",), sz=8, color=DIVIDER_HEX)
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            cell = table.rows[i].cells[j]
            _cell_text(cell, str(val), size=SIZE_BODY, color=BODY_SOFT)
            cell.width = Inches(col_widths[j])
            _border(cell, sides=("top", "bottom"), sz=4, color=LIGHT_BORDER_HEX)
    _add_paragraph(doc, "", space_after=4)


def render_strategy_definition(doc, sd: dict) -> None:
    # The content section was started by render_process_flow, so just break
    # to a new page rather than starting a new section.
    _page_break(doc)
    _add_heading(doc, "Strategy Definition", 1)

    _add_heading(doc, "Idea Description", 2)
    idea = sd.get("idea_description", {})
    if idea.get("paragraph"):
        _add_paragraph(doc, idea["paragraph"])
    for b in idea.get("mechanisms", []) or []:
        _add_bullet(doc, b)
    if idea.get("closing_paragraph"):
        _add_paragraph(doc, idea["closing_paragraph"])

    _add_heading(doc, "Background Research", 2)
    _add_paragraph(doc,
                   sd.get("background_research_intro",
                          "Momentum and trend-following strategies have been "
                          "extensively documented in academic financial research."))
    for paper in sd.get("background_research", []) or []:
        _add_heading(doc, paper.get("title", ""), 3)
        if paper.get("date"):
            para = doc.add_paragraph()
            r = para.add_run("Date: ")
            _set_run(r, size=SIZE_BODY, bold=True)
            r2 = para.add_run(paper["date"])
            _set_run(r2, size=SIZE_BODY)
        if paper.get("authors"):
            para = doc.add_paragraph()
            r = para.add_run("Authors: ")
            _set_run(r, size=SIZE_BODY, bold=True)
            r2 = para.add_run(", ".join(paper["authors"]))
            _set_run(r2, size=SIZE_BODY)
        if paper.get("link"):
            para = doc.add_paragraph()
            r = para.add_run("Link: ")
            _set_run(r, size=SIZE_BODY, bold=True)
            r2 = para.add_run(paper["link"])
            _set_run(r2, size=SIZE_BODY,
                     color=RGBColor(0x10, 0x71, 0xE5))
            r2.font.underline = True
        for b in paper.get("findings", []) or []:
            _add_bullet(doc, b)

    _add_heading(doc, "Investable Universe and Constraints", 2)
    _add_heading(doc, "Universe", 3)
    for b in sd.get("universe", []) or []:
        _add_bullet(doc, b)
    _add_heading(doc, "Strategy Constraints", 3)
    for b in sd.get("constraints", []) or []:
        _add_bullet(doc, b)

    _add_heading(doc, "Benchmark", 2)
    bench = sd.get("benchmark", {})
    if bench.get("primary"):
        _add_bullet(doc, bench["primary"])
    if bench.get("secondary"):
        s = bench["secondary"]
        _add_bullet(doc, s.get("name", ""))
        for key, label in [("fmp_ticker", "FMP Ticker"),
                           ("isin", "ISIN"),
                           ("ipo", "IPO"),
                           ("purpose", "Used as")]:
            if s.get(key):
                _add_bullet(doc, f"{label}: {s[key]}", level=1)
    if bench.get("closing_paragraph"):
        _add_paragraph(doc, bench["closing_paragraph"])

    _add_heading(doc, "Hypothesis", 2)
    hyp = sd.get("hypothesis", {})
    _add_heading(doc, "Market Inefficiency", 3)
    for b in hyp.get("market_inefficiency", []) or []:
        _add_bullet(doc, b)
    _add_heading(doc, "Behavioral", 3)
    for b in hyp.get("behavioral", []) or []:
        _add_bullet(doc, b)
    _add_heading(doc, "Structural", 3)
    for b in hyp.get("structural", []) or []:
        _add_bullet(doc, b)
    _add_heading(doc, "Null Hypothesis", 3)
    if hyp.get("null"):
        _add_bullet(doc, hyp["null"])

    _add_heading(doc, "Unified Modeling Language Diagrams", 2)
    uml = sd.get("uml_diagram", {})
    _image_or_placeholder(doc, uml.get("image_path"),
                          uml.get("caption", "Strategy UML Diagram"))


def render_feature_engineering(doc, fe: dict) -> None:
    _page_break(doc)
    _add_heading(doc, "Feature Engineering", 1)

    data = fe.get("data", {})
    _add_heading(doc, "Data", 2)
    _add_paragraph(doc, data.get("intro",
                                 "The strategy relies primarily on market price data."))
    _add_heading(doc, "Key Variables", 3)
    for b in data.get("key_variables", []) or []:
        _add_bullet(doc, b)
    _add_heading(doc, "Adjusted Prices Incorporate", 3)
    for b in data.get("adjusted_prices_incorporate", []) or []:
        _add_bullet(doc, b)
    if data.get("adjustment_note"):
        _add_paragraph(doc, data["adjustment_note"])

    da = fe.get("data_analysis", {})
    _add_heading(doc, "Data Analysis", 2)
    _add_heading(doc, "Time Series and Cross-Sectional Analysis", 3)
    for b in da.get("time_series_and_cross_sectional", []) or []:
        _add_bullet(doc, b)
    _add_heading(doc, "Key Financial Metrics", 3)
    for b in da.get("key_financial_metrics", []) or []:
        _add_bullet(doc, b)
    _add_heading(doc, "Common Visualizations", 3)
    for b in da.get("common_visualizations", []) or []:
        _add_bullet(doc, b)

    feats = fe.get("features", {})
    _add_heading(doc, "Features", 2)
    if feats.get("intro_paragraph"):
        _add_paragraph(doc, feats["intro_paragraph"])
    for grp in feats.get("groups", []) or []:
        _add_heading(doc, grp.get("title", "Features"), 3)
        rows = [[r.get("feature", ""), r.get("description", "")]
                for r in grp.get("rows", [])]
        _generic_table(doc, ["Feature", "Description"], rows,
                       col_widths=[1.5, 5.0])
        if grp.get("note"):
            para = doc.add_paragraph()
            r = para.add_run(f"Note: {grp['note']}")
            _set_run(r, size=SIZE_BODY)
            r.italic = True
    if feats.get("overview_image"):
        _image_or_placeholder(doc, feats["overview_image"], "Features Overview")

    _add_heading(doc, "Strategy Signals", 2)
    for sig in fe.get("strategy_signals", []) or []:
        _add_heading(doc, sig.get("name", "Signal"), 3)
        if sig.get("intro"):
            _add_paragraph(doc, sig["intro"])
        for line in sig.get("branches", []) or []:
            _add_paragraph(doc, line, size=SIZE_SIGNAL, space_after=2)
        _add_heading(doc, "Signal Logic", 4)
        for b in sig.get("logic", []) or []:
            _add_paragraph(doc, f"• {b}")


def render_strategy_design(doc, sdes: dict) -> None:
    _page_break(doc)
    _add_heading(doc, "Strategy Design", 1)

    sm = sdes.get("strategy_modeling", {})
    _add_heading(doc, "Strategy Modeling", 2)
    for para in sm.get("intro_paragraphs", []) or []:
        _add_paragraph(doc, para)
    rows = [[r.get("dimension", ""), r.get("question", "")]
            for r in sm.get("table", []) or []]
    _generic_table(doc, ["Dimension", "Question"], rows,
                   col_widths=[1.4, 5.1])

    pc = sdes.get("portfolio_construction", {})
    _add_heading(doc, "Portfolio Construction", 2)
    _add_heading(doc, "Selection Rules", 3)
    for b in pc.get("selection_rules", []) or []:
        _add_paragraph(doc, f"• {b}")
    _add_heading(doc, "Sizing Rules", 3)
    for b in pc.get("sizing_rules", []) or []:
        _add_paragraph(doc, f"• {b}")
    _add_heading(doc, "Rebalancing Logic — Event-Driven", 3)
    rb = pc.get("rebalancing_logic", {})
    if rb.get("intro"):
        _add_paragraph(doc, rb["intro"])
    for b in rb.get("bullets", []) or []:
        _add_paragraph(doc, f"• {b}")
    _add_heading(doc, "Market Regime Handling", 3)
    if pc.get("market_regime_handling"):
        _add_paragraph(doc, pc["market_regime_handling"])

    bt = sdes.get("backtesting", {})
    _add_heading(doc, "Backtesting", 2)
    _add_heading(doc, "Backtesting Config", 3)
    cfg_rows = [[r.get("parameter", ""), r.get("value", ""), r.get("description", "")]
                for r in bt.get("config", []) or []]
    if cfg_rows:
        _generic_table(doc, ["Parameter", "Value", "Description"], cfg_rows,
                       col_widths=[1.8, 1.7, 3.0])
    for sub_title, key in [
        ("Performance Top 20", "performance_image"),
        ("Commissions Top 20", "commissions_image"),
        ("Drawdown Top 20", "drawdown_image"),
        ("Annual Returns Top 20", "annual_returns_image"),
    ]:
        _add_heading(doc, sub_title, 3)
        _image_or_placeholder(doc, bt.get(key), sub_title)

    _add_heading(doc, "Attribution Analysis", 2)
    attr = sdes.get("attribution_analysis", {})
    if attr.get("image_path"):
        _image_or_placeholder(doc, attr["image_path"], "Attribution Analysis")
    for i, finding in enumerate(attr.get("findings", []) or [], start=1):
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.25)
        para.paragraph_format.space_after = Pt(4)
        r = para.add_run(f"{i}. ")
        _set_run(r, size=SIZE_BODY)
        r2 = para.add_run(finding.get("headline", ""))
        _set_run(r2, size=SIZE_BODY, bold=True)
        r3 = para.add_run(" " + finding.get("body", ""))
        _set_run(r3, size=SIZE_BODY)

    _add_heading(doc, "Strategy Score", 2)
    score = sdes.get("strategy_score", {})
    if score.get("rows"):
        _render_strategy_score(doc, score)
    for i, obs in enumerate(score.get("observations", []) or [], start=1):
        _add_paragraph(doc, f"{i}. {obs}", left_indent=0.25)


def _render_strategy_score(doc, score: dict) -> None:
    rows_in = score.get("rows", [])
    headers = ["Strategy Blueprint", "Scoring Concept", "Weight",
               "Score (0-100)", "Weighted Score", "Gap (Score)"]
    table = doc.add_table(rows=2 + len(rows_in), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    col_widths = [1.4, 2.0, 0.7, 0.85, 1.0, 0.85]

    # header
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        _cell_text(cell, h, size=SIZE_BODY, bold=True, color=BODY,
                   align=(WD_ALIGN_PARAGRAPH.RIGHT if j >= 2
                          else WD_ALIGN_PARAGRAPH.LEFT))
        cell.width = Inches(col_widths[j])
        _border(cell, sides=("bottom",), sz=12, color="000000")

    weighted_values = []
    total_weight = 0.0
    total_weighted = 0.0
    last_section = None
    for i, row in enumerate(rows_in, start=1):
        section = row.get("section", "")
        concept = row.get("concept", "")
        weight = float(row.get("weight", 0))
        score_v = row.get("score", "")
        if isinstance(score_v, (int, float)):
            weighted = weight * float(score_v) / 100.0
            gap = weight - weighted
            weighted_values.append(weighted)
            total_weighted += weighted
            score_cell_str = f"{int(score_v)}"
            weighted_cell_str = f"{weighted:.1f}"
            gap_cell_str = f"{gap:.1f}"
        else:
            score_cell_str = "TODO"
            weighted_cell_str = ""
            gap_cell_str = ""
        total_weight += weight
        section_label = section if section != last_section else ""
        last_section = section

        cells = table.rows[i].cells
        _cell_text(cells[0], section_label, size=SIZE_BODY, bold=bool(section_label))
        _cell_text(cells[1], concept, size=SIZE_BODY, color=BODY_SOFT)
        _cell_text(cells[2], f"{weight:.2f}%", size=SIZE_BODY, color=BODY_SOFT,
                   align=WD_ALIGN_PARAGRAPH.RIGHT)
        _cell_text(cells[3], score_cell_str, size=SIZE_BODY, color=BODY_SOFT,
                   align=WD_ALIGN_PARAGRAPH.RIGHT)
        _cell_text(cells[4], weighted_cell_str, size=SIZE_BODY, color=BODY_SOFT,
                   align=WD_ALIGN_PARAGRAPH.RIGHT)
        _cell_text(cells[5], gap_cell_str, size=SIZE_BODY, color=BODY_SOFT,
                   align=WD_ALIGN_PARAGRAPH.RIGHT)
        for j in range(len(headers)):
            cells[j].width = Inches(col_widths[j])
            _border(cells[j], sides=("top", "bottom"), sz=4,
                    color=LIGHT_BORDER_HEX)

    # Apply gradient to Weighted Score column
    if weighted_values:
        lo, hi = min(weighted_values), max(weighted_values)
        for i, w in enumerate(weighted_values, start=1):
            cell = table.rows[i].cells[4]
            _shade(cell, _interpolate(w, lo, hi))

    # Coral on Gap > 0
    for i, row in enumerate(rows_in, start=1):
        score_v = row.get("score", None)
        if isinstance(score_v, (int, float)):
            weight = float(row.get("weight", 0))
            gap = weight - (weight * float(score_v) / 100.0)
            if gap > 1e-6:
                cell = table.rows[i].cells[5]
                _shade(cell, SCORE_LOW_HEX)
                # Override text colour to white
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Overall row
    total_row = table.rows[-1]
    _cell_text(total_row.cells[0], "Overall Score", size=SIZE_BODY, bold=True)
    _cell_text(total_row.cells[1], "")
    _cell_text(total_row.cells[2], f"{total_weight:.2f}%", size=SIZE_BODY,
               bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _cell_text(total_row.cells[3], "")
    _cell_text(total_row.cells[4], f"{total_weighted:.1f}", size=SIZE_BODY,
               bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _cell_text(total_row.cells[5], "")
    for j in range(len(headers)):
        total_row.cells[j].width = Inches(col_widths[j])
        _border(total_row.cells[j], sides=("top",), sz=8, color="000000")

    _add_paragraph(doc, "", space_after=4)


def render_conclusions(doc, cc: dict) -> None:
    _page_break(doc)
    _add_heading(doc, "Conclusions", 1)

    hv = cc.get("hypothesis_validation", {})
    _add_heading(doc, "Hypothesis Validation.", 2)
    if hv.get("lead_question"):
        _add_paragraph(doc, hv["lead_question"], bold=True)
    if hv.get("key_findings"):
        _add_paragraph(doc, "Key Findings:", bold=True)
        for b in hv["key_findings"]:
            _add_bullet(doc, b)
    if hv.get("interpretation"):
        _add_paragraph(doc, "Interpretation:", bold=True)
        for b in hv["interpretation"]:
            _add_bullet(doc, b)
    if hv.get("conclusion"):
        _add_paragraph(doc, "Conclusion:", bold=True)
        _add_bullet(doc, hv["conclusion"])

    impl = cc.get("implementable", {})
    _add_heading(doc, "Is This An Implementable Strategy?", 2)
    if impl.get("implementation_risks"):
        _add_paragraph(doc, "Implementation Risks", bold=True)
        for risk in impl["implementation_risks"]:
            _add_paragraph(doc, risk.get("name", ""), bold=True, space_after=2)
            _add_arrow(doc, risk.get("explanation", ""))
    if impl.get("operational_requirements"):
        _add_paragraph(doc, "Operational Requirements", bold=True)
        for line in impl["operational_requirements"]:
            _add_arrow(doc, line)
    if impl.get("decision"):
        _add_paragraph(doc, "Decision", bold=True)
        _add_paragraph(doc, f"\U0001F449 {impl['decision']}")
    if impl.get("key_condition"):
        _add_paragraph(doc, "Key Condition", bold=True)
        _add_paragraph(doc, impl["key_condition"])

    _add_heading(doc, "Next Steps", 2)
    for i, group in enumerate(cc.get("next_steps", []) or [], start=1):
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.25)
        para.paragraph_format.space_after = Pt(4)
        r = para.add_run(f"{i}. ")
        _set_run(r, size=SIZE_BODY)
        r2 = para.add_run(group.get("category", ""))
        _set_run(r2, size=SIZE_BODY, bold=True)
        for item in group.get("items", []) or []:
            _add_bullet(doc, item, level=1)


def render_findings(doc, ff: dict) -> None:
    _page_break(doc)
    _add_heading(doc, "Findings, Concerns and Decisions", 1)
    for it in ff.get("iterations", []) or []:
        _add_heading(doc, it.get("title", "Iteration"), 2)
        for entry in it.get("entries", []) or []:
            _add_heading(doc, entry.get("date", ""), 3)
            _add_paragraph(doc, entry.get("body", ""))


# ---------------------------------------------------------------------------
# Style setup — set the default font and heading colours on the document
# ---------------------------------------------------------------------------

def _configure_styles(doc) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = FONT_REG
    normal.font.size = Pt(SIZE_BODY)
    rPr = normal.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), FONT_REG)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def build(content_path: str, output_path: str,
          strategy_name: Optional[str]) -> None:
    with open(content_path, "r", encoding="utf-8") as fh:
        content = json.load(fh)
    if strategy_name:
        content["strategy_name"] = strategy_name

    score_rows = (
        content.get("strategy_design", {})
        .get("strategy_score", {})
        .get("rows", [])
    )
    if score_rows:
        total = sum(float(r.get("weight", 0)) for r in score_rows)
        if abs(total - 100.0) > 0.05:
            raise ValueError(
                f"strategy_score.rows[*].weight must sum to 100.0 (got {total})"
            )

    doc = Document()
    _configure_styles(doc)
    render_cover(doc, content)
    render_process_flow(doc, content.get("process_flow", {
        "intro": "The Process Flow:",
        "items": [
            {"label": "Strategy Definition",
             "description": "Moving from intuition to a clear, testable hypothesis with an explicit source of returns."},
            {"label": "Feature Engineering",
             "description": "Transforming raw data into meaningful, hypothesis-driven signals."},
            {"label": "Strategy Design",
             "description": "Translating signals into portfolio decisions — selection, sizing, and timing — and validating results through robust testing and analysis."},
        ],
    }))
    render_strategy_definition(doc, content.get("strategy_definition", {}))
    render_feature_engineering(doc, content.get("feature_engineering", {}))
    render_strategy_design(doc, content.get("strategy_design", {}))
    render_conclusions(doc, content.get("conclusions", {}))
    render_findings(doc, content.get("findings", {}))

    doc.save(output_path)
    print(f"Wrote {output_path}")


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--content", required=True, help="Path to content.json")
    parser.add_argument("--output", required=True, help="Output .docx path")
    parser.add_argument("--strategy-name", default=None,
                        help="Override strategy name (otherwise read from content.json)")
    args = parser.parse_args()
    build(args.content, args.output, args.strategy_name)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
